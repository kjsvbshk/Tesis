"""
Reemplaza los diagramas de schema en manual_tecnico_haw.docx.
Elimina las imágenes de la sección 3.2 ya insertadas y coloca las nuevas.
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(__file__).parent.parent
DOCS = Path(__file__).parent
DOCX = ROOT / "manual_tecnico_haw.docx"
IMG_ESPN = DOCS / "schema_espn.png"
IMG_APP  = DOCS / "schema_app.png"
IMG_ML   = DOCS / "schema_ml.png"


def has_picture(para):
    """Retorna True si el párrafo contiene una imagen inline."""
    return bool(para._element.findall('.//' + qn('a:blip')))


def is_caption(para):
    text = para.text.strip()
    return text.startswith("Figura 2")


def find_section_32_heading(doc):
    for p in doc.paragraphs:
        if "3.2" in p.text and "Diagrama" in p.text:
            return p
    return None


def remove_existing_images_and_captions(doc, after_para):
    """Elimina párrafos de imagen y caption Figura 2x después del heading."""
    body = doc.element.body
    elems = list(body)
    start_idx = elems.index(after_para._element)

    to_remove = []
    # Scan next 20 elements for pictures or Figura 2 captions
    for i in range(start_idx + 1, min(start_idx + 25, len(elems))):
        elem = elems[i]
        # Check if this element is a paragraph
        if elem.tag.endswith('}p'):
            from docx.text.paragraph import Paragraph
            p = Paragraph(elem, body)
            if has_picture(p):
                to_remove.append(elem)
            elif is_caption(p):
                to_remove.append(elem)

    for elem in to_remove:
        body.remove(elem)
    print(f"  Eliminados {len(to_remove)} elementos (imágenes + captions) de sección 3.2")


def insert_image_after(doc, anchor_para, img_path, width_inches, caption_text):
    body = doc.element.body
    anchor_idx = list(body).index(anchor_para._element)

    # Imagen
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(str(img_path), width=Inches(width_inches))
    body.remove(img_p._element)
    body.insert(anchor_idx + 1, img_p._element)

    # Caption
    cap_p = doc.add_paragraph(caption_text, style="Normal")
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap_p.runs:
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    body.remove(cap_p._element)
    body.insert(anchor_idx + 2, cap_p._element)

    return cap_p


def main():
    print(f"Abriendo: {DOCX}")
    doc = Document(str(DOCX))

    heading = find_section_32_heading(doc)
    if heading is None:
        print("[ERROR] No se encontró el heading de sección 3.2")
        return

    print(f"  Heading encontrado: '{heading.text}'")

    # Eliminar imágenes y captions anteriores
    remove_existing_images_and_captions(doc, heading)

    # Insertar las 3 nuevas imágenes
    cap1 = insert_image_after(
        doc, heading, IMG_ESPN, 6.2,
        "Figura 2a. Diagrama del schema ESPN — tablas de scraping y cuotas de apuestas."
    )
    cap2 = insert_image_after(
        doc, cap1, IMG_APP, 6.2,
        "Figura 2b. Diagrama del schema APP — usuarios, apuestas, predicciones, seguridad y auditoría."
    )
    insert_image_after(
        doc, cap2, IMG_ML, 5.0,
        "Figura 2c. Diagrama del schema ML — tabla ml_ready_games con features de entrenamiento."
    )

    doc.save(str(DOCX))
    print(f"[OK] Guardado: {DOCX}")


if __name__ == "__main__":
    main()
