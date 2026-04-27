"""
Script de actualización del manual_tecnico_haw.docx
Actualiza secciones de BD, ML y Scrapping con el contexto actual del proyecto.
Inserta los diagramas de schema generados en docs/.
"""

import os
import sys
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent.parent
DOCS = Path(__file__).parent
DOCX_IN  = ROOT / "manual_tecnico_haw.docx"
DOCX_OUT = ROOT / "manual_tecnico_haw.docx"   # sobreescribe el original

IMG_ESPN = DOCS / "schema_espn.png"
IMG_APP  = DOCS / "schema_app.png"
IMG_ML   = DOCS / "schema_ml.png"

# ── Helpers ────────────────────────────────────────────────────────

def clear_paragraph(p):
    """Elimina todo el contenido de un párrafo."""
    for run in p.runs:
        run.text = ""
    p.clear()

def set_paragraph_text(p, text, bold=False, size=None):
    """Reemplaza el texto de un párrafo preservando el estilo."""
    p.clear()
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)

def add_paragraph_after(p, text, style="Normal"):
    """Inserta un nuevo párrafo después del párrafo dado."""
    new_p = OxmlElement("w:p")
    p._element.addnext(new_p)
    new_para = p._element.getnext()
    # Usamos el API de Document para asignar estilo
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, p._element.getparent())
    para.style = style
    para.add_run(text)
    return para

def insert_image_after(doc, anchor_paragraph, img_path, width_inches=6.2, caption=None):
    """Inserta una imagen después de anchor_paragraph."""
    # Encontrar índice del párrafo anchor en el body
    body = doc.element.body
    paras = list(body)
    anchor_elem = anchor_paragraph._element
    idx = list(body).index(anchor_elem)

    # Crear párrafo de imagen
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))

    # Mover el párrafo al lugar correcto
    body.remove(img_para._element)
    body.insert(idx + 1, img_para._element)

    # Crear párrafo de caption
    if caption:
        cap_para = doc.add_paragraph(caption, style="Normal")
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap_para.runs:
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        body.remove(cap_para._element)
        body.insert(idx + 2, cap_para._element)
        return img_para, cap_para

    return img_para, None


def find_paragraph_by_text(doc, search_text, heading_level=None):
    """Busca un párrafo por texto parcial y estilo opcional."""
    for p in doc.paragraphs:
        if search_text.lower() in p.text.lower():
            if heading_level is None:
                return p
            if f"Heading {heading_level}" in p.style.name or f"heading {heading_level}" in p.style.name.lower():
                return p
    return None


def find_paragraph_by_index(doc, idx):
    return doc.paragraphs[idx]


def insert_paragraph_after_element(doc, ref_elem, text, style="Normal", bold=False, italic=False, size=None):
    """Inserta un párrafo en el body justo después de ref_elem."""
    new_p = doc.add_paragraph(text, style=style)
    for run in new_p.runs:
        if bold: run.bold = True
        if italic: run.font.italic = True
        if size: run.font.size = Pt(size)
    # Mover al lugar correcto
    body = doc.element.body
    body.remove(new_p._element)
    ref_idx = list(body).index(ref_elem)
    body.insert(ref_idx + 1, new_p._element)
    return new_p


# ── Actualización principal ─────────────────────────────────────────

def update_manual():
    print(f"Abriendo: {DOCX_IN}")
    doc = Document(str(DOCX_IN))
    paras = doc.paragraphs

    # ──────────────────────────────────────────────────────────────
    # 1. SECCIÓN 3 — Modelo de Base de Datos
    # ──────────────────────────────────────────────────────────────

    # 3.0 Párrafo introductorio (idx 67)
    p_db_intro = find_paragraph_by_text(doc, "La base de datos está organizada en cuatro schemas")
    if p_db_intro is None:
        p_db_intro = find_paragraph_by_text(doc, "cuatro schemas independientes")
    if p_db_intro:
        set_paragraph_text(p_db_intro,
            "El sistema utiliza PostgreSQL alojado en Neon (servicio cloud serverless). "
            "La base de datos está organizada en tres schemas independientes que separan los datos "
            "según su dominio y responsable: espn (datos externos de ESPN gestionados por el módulo Scrapping), "
            "ml (features engineered y dataset de entrenamiento gestionado por el módulo ML), "
            "y app (datos de la aplicación: usuarios, apuestas, predicciones, auditoría, RBAC, "
            "sesiones, idempotencia y versionado de modelos)."
        )
        print("[OK] Párrafo intro sección 3 actualizado")

    # 3.2 — Eliminar fig caption obsoleto e insertar los 3 diagramas
    p_322 = find_paragraph_by_text(doc, "3.2. Diagrama", heading_level=2)
    if p_322:
        # Buscar el caption antiguo (Figura 2) y vaciarlo
        p_fig2 = find_paragraph_by_text(doc, "Figura 2.")
        if p_fig2:
            set_paragraph_text(p_fig2, "")

        # Insertar los 3 imágenes después del heading 3.2
        last_inserted = p_322
        _, cap1 = insert_image_after(
            doc, last_inserted, IMG_ESPN, width_inches=6.2,
            caption="Figura 2a. Diagrama del schema ESPN — tablas de scraping y cuotas de apuestas."
        )
        last_ref = cap1 if cap1 else last_inserted
        _, cap2 = insert_image_after(
            doc, last_ref, IMG_APP, width_inches=6.2,
            caption="Figura 2b. Diagrama del schema APP — usuarios, apuestas, predicciones, seguridad y auditoría."
        )
        last_ref2 = cap2 if cap2 else last_ref
        _, cap3 = insert_image_after(
            doc, last_ref2, IMG_ML, width_inches=5.0,
            caption="Figura 2c. Diagrama del schema ML — tabla ml_ready_games con features de entrenamiento."
        )
        print("[OK] Imágenes de schemas insertadas en sección 3.2")

    # 3.3 — Renombrar "schema sys" → "schema app" y actualizar tablas
    p_33 = find_paragraph_by_text(doc, "3.3. Tablas principales", heading_level=2)
    if p_33:
        set_paragraph_text(p_33, "3.3. Tablas principales — schema app (aplicación)", bold=False)
        print("[OK] Heading 3.3 actualizado")

    # 3.3.1 users → user_accounts
    p_331 = find_paragraph_by_text(doc, "3.3.1.", heading_level=3)
    if p_331:
        set_paragraph_text(p_331, "3.3.1. user_accounts, clients, administrators, operators")
        print("[OK] Heading 3.3.1 actualizado")

    p_users_text = find_paragraph_by_text(doc, "Almacena los datos de registro y autenticación de todos los usuarios")
    if p_users_text:
        set_paragraph_text(p_users_text,
            "user_accounts es la tabla base de autenticación (username, email, hashed_password con bcrypt, "
            "is_active). El perfil específico se almacena en tablas derivadas: clients (usuarios con créditos "
            "virtuales y datos personales), administrators (con employee_id y department) y operators "
            "(con turno de trabajo). Este diseño polimórfico facilita la extensión de tipos de usuario "
            "sin modificar la tabla base."
        )
        print("[OK] Párrafo user_accounts actualizado")

    # 3.3.2 roles
    p_332 = find_paragraph_by_text(doc, "3.3.2.", heading_level=3)
    if p_332:
        set_paragraph_text(p_332, "3.3.2. roles, permissions y tablas de asociación RBAC")
        print("[OK] Heading 3.3.2 actualizado")

    p_roles_text = find_paragraph_by_text(doc, "Implementan el control de acceso basado en roles (RBAC)")
    if p_roles_text:
        set_paragraph_text(p_roles_text,
            "Implementan el control de acceso basado en roles (RBAC). La tabla roles define los roles "
            "disponibles (cliente, administrador, operador). permissions contiene permisos granulares por "
            "scope (ej. predictions:read, bets:write, admin:all). Las tablas de asociación user_roles y "
            "role_permissions establecen las relaciones N:M. La autorización se verifica en cada request "
            "mediante el servicio authorization.py."
        )
        print("[OK] Párrafo roles actualizado")

    # 3.3.3 audit_log
    p_333 = find_paragraph_by_text(doc, "3.3.3.", heading_level=3)
    if p_333:
        set_paragraph_text(p_333, "3.3.3. audit_log, outbox y requests")
        print("[OK] Heading 3.3.3 actualizado")

    p_audit_text = find_paragraph_by_text(doc, "Registra todas las acciones significativas")
    if p_audit_text:
        set_paragraph_text(p_audit_text,
            "audit_log registra todas las acciones significativas del sistema: autenticaciones, apuestas, "
            "cambios de perfil, acciones de administración. Incluye actor_user_id, action, resource_type, "
            "resource_id y snapshots JSON before/after para trazabilidad completa. "
            "outbox implementa el Outbox Pattern: los eventos (bet.placed, prediction.completed) se "
            "insertan en la misma transacción DB que la operación principal; un worker los procesa "
            "asíncronamente. published_at NULL indica evento pendiente. "
            "requests realiza el tracking del ciclo de vida de cada request (RECEIVED → PROCESSING → "
            "COMPLETED / FAILED) con request_key para idempotencia."
        )
        print("[OK] Párrafo audit/outbox actualizado")

    # 3.4 — "schema app" → "schema espn"
    p_34 = find_paragraph_by_text(doc, "3.4. Tablas principales", heading_level=2)
    if p_34:
        set_paragraph_text(p_34, "3.4. Tablas principales — schema espn (datos de scraping)")
        print("[OK] Heading 3.4 actualizado")

    # 3.4.1 bets → games
    p_341 = find_paragraph_by_text(doc, "3.4.1.", heading_level=3)
    if p_341:
        set_paragraph_text(p_341, "3.4.1. games y teams")
        print("[OK] Heading 3.4.1 actualizado")

    p_bets_old = find_paragraph_by_text(doc, "La tabla bets registra cada apuesta")
    if not p_bets_old:
        # Buscar el párrafo después de 3.4.1 heading
        p_bets_old = find_paragraph_by_text(doc, "monto apostado")
    if p_bets_old:
        set_paragraph_text(p_bets_old,
            "games es la tabla central del schema espn. Almacena un registro por partido NBA con: "
            "game_id (PK BigInt de ESPN), fecha, home_team/away_team (nombres normalizados), "
            "scores, estadísticas agregadas de box (FG%, 3P%, FT%, rebotes, asistencias, etc.) y "
            "variables derivadas (home_win, point_diff, net_rating_diff). "
            "teams almacena el catálogo de los 30 equipos NBA con abbreviation, conference y division. "
            "team_stats_game complementa games con estadísticas detalladas por equipo y partido."
        )
        print("[OK] Párrafo games/teams actualizado")

    # 3.4.2 predictions_cache → bets/odds
    p_342 = find_paragraph_by_text(doc, "3.4.2.", heading_level=3)
    if p_342:
        set_paragraph_text(p_342, "3.4.2. bets, bet_selections, bet_results, game_odds")
        print("[OK] Heading 3.4.2 actualizado")

    p_pred_cache = find_paragraph_by_text(doc, "Almacena las predicciones generadas por el modelo ML con un TTL")
    if p_pred_cache:
        set_paragraph_text(p_pred_cache,
            "bets registra cada apuesta con su tipo (moneyline, spread, over/under), estado "
            "(pending, won, lost, cancelled), monto y cuota en el momento de la apuesta (snapshot de auditoría). "
            "bet_selections detalla la selección específica (equipo elegido, valor de spread, "
            "valor de over/under). bet_results almacena el resultado final y pago efectivo. "
            "game_odds almacena las cuotas de mercado por partido, tipo y proveedor con snapshots temporales."
        )
        print("[OK] Párrafo bets/odds actualizado")

    # 3.5 — Configuración conexión → schema ml
    p_35 = find_paragraph_by_text(doc, "3.5.", heading_level=2)
    if p_35:
        set_paragraph_text(p_35, "3.5. Tablas principales — schema ml (machine learning)")
        print("[OK] Heading 3.5 actualizado")

    # Buscar párrafo de configuración de conexión y actualizarlo
    p_conn = find_paragraph_by_text(doc, "La conexión a Neon se configura")
    if not p_conn:
        p_conn = find_paragraph_by_text(doc, "NEON_DB_HOST")
    if p_conn:
        set_paragraph_text(p_conn,
            "ml_ready_games es la única tabla del schema ml. Contiene una fila por partido NBA con todas "
            "las features engineered listas para entrenar. La columna home_win es el target binario "
            "(True = victoria local). Las features se dividen en: rolling statistics con ventanas de 5 y "
            "10 juegos (ppg, net_rating, win_rate, efg_pct, tov_rate, reb, ast, tov), factores "
            "contextuales (rest_days, back-to-back, injury_count, Elo, streak) y diferenciales "
            "(home_value − away_value para cada feature). Las features marcadas v2.0.0 son nuevas en "
            "ese modelo candidato. Cobertura: ~1,700 partidos de temporadas 2023-24 a 2025-26."
        )
        print("[OK] Párrafo schema ML actualizado")

    # ──────────────────────────────────────────────────────────────
    # 2. SECCIÓN 5 — Módulo de Machine Learning
    # ──────────────────────────────────────────────────────────────

    # 5.0 Párrafo introductorio
    p_ml_intro = find_paragraph_by_text(doc, "El módulo ML es responsable de preparar el dataset")
    if p_ml_intro:
        set_paragraph_text(p_ml_intro,
            "El módulo ML es responsable de preparar el dataset de entrenamiento, entrenar los modelos "
            "de predicción y exportarlos en formato .joblib para que el backend pueda cargarlos en tiempo "
            "de ejecución. El pipeline completo ha sido ejecutado hasta la Fase 4 (entrenamiento). "
            "El modelo activo en producción es v1.6.0 (Ensemble RF+XGBoost con calibración Isotonic), "
            "que pasa todos los criterios de aceptación. El modelo v2.0.0 (33 features, incluyendo Elo, "
            "H2H y splits locales/visitantes) fue entrenado como candidato pero aún no está integrado "
            "en el Backend."
        )
        print("[OK] Párrafo intro sección 5 actualizado")

    # 5.1 Feature engineering
    p_51_text = find_paragraph_by_text(doc, "El feature engineering es el proceso de transformar los datos crudos")
    if p_51_text:
        set_paragraph_text(p_51_text,
            "El feature engineering transforma los datos crudos de ESPN en variables predictoras. "
            "El script principal src/etl/build_features.py es idempotente y calcula todas las features "
            "usando exclusivamente datos de partidos anteriores a la fecha del juego (shift(1)), "
            "previniendo data leakage. Las features del modelo v1.6.0 (21 features) incluyen: "
            "rolling statistics por equipo en ventanas de 5 y 10 juegos (ppg, net_rating, win_rate), "
            "días de descanso (rest_days), indicador de back-to-back (b2b) y conteo de lesiones activas. "
            "El modelo v2.0.0 amplía a 33 features añadiendo EFG%, Turnover Rate, OReb%, DReb%, "
            "Elo rating, rachas de victorias/derrotas y ventaja histórica H2H."
        )
        print("[OK] Párrafo 5.1 feature engineering actualizado")

    # 5.2 Modelos entrenados
    p_52_text = find_paragraph_by_text(doc, "Se entrenan y evalúan dos algoritmos de clasificación")
    if not p_52_text:
        p_52_text = find_paragraph_by_text(doc, "dos algoritmos de clasificaci")
    if p_52_text:
        set_paragraph_text(p_52_text,
            "El sistema entrena cinco componentes organizados en un pipeline de ensamble:\n"
            "• NBARandomForestModel (src/models/random_forest.py): clasificador binario P(home_win).\n"
            "• NBAXGBoostModel (src/models/xgboost_model.py): clasificador binario P(home_win).\n"
            "• NBAEnsembleModel (src/models/ensemble.py): stacking de RF+XGBoost con meta-modelo "
            "LogisticRegression y calibración Isotonic Regression para garantizar probabilidades realistas.\n"
            "• NBAMarginModel (src/models/margin_model.py): regresor XGBoost para el margen esperado "
            "de puntos (home_score − away_score).\n"
            "• NBATotalModel (src/models/total_model.py): regresor XGBoost para el total esperado "
            "de puntos del partido.\n\n"
            "Métricas del modelo activo v1.6.0 (test set, n=734): Log Loss 0.6553, Brier Score 0.2312, "
            "ROC-AUC 0.6542, ECE 0.0363, Accuracy 62.1%. Todos los criterios de aceptación superados."
        )
        print("[OK] Párrafo 5.2 modelos actualizado")

    # 5.3 Pipeline de entrenamiento
    p_53 = find_paragraph_by_text(doc, "5.3. Pipeline de entrenamiento", heading_level=2)
    if p_53:
        # Buscar el párrafo de código/comandos que sigue
        p_53_body = find_paragraph_by_text(doc, "train_model")
        if not p_53_body:
            p_53_body = find_paragraph_by_text(doc, "pipeline de entrenamiento ejecuta")
        if p_53_body:
            set_paragraph_text(p_53_body,
                "El pipeline de entrenamiento se ejecuta desde src/training/train.py. Los pasos son: "
                "(1) carga del dataset ml.ml_ready_games desde Neon, (2) split temporal train/test "
                "(sin shuffle para respetar la secuencia temporal), (3) entrenamiento de RF y XGBoost "
                "por separado, (4) calibración Isotonic del ensamble, (5) entrenamiento de MarginModel "
                "y TotalModel, (6) evaluación con metrics.py (Log Loss, Brier, ROC-AUC, ECE, MAE), "
                "(7) exportación del modelo a ML/models/ y copia a Backend/ml/models/ mediante "
                "scripts/deploy_model.py. Los resultados se guardan en reports/ con backtesting "
                "y comparación contra baselines (always_home, random)."
            )
            print("[OK] Párrafo 5.3 pipeline actualizado")

    # 5.4 Versionado de modelos
    p_54_text = find_paragraph_by_text(doc, "Cada modelo exportado se registra en la tabla sys.model_versions")
    if not p_54_text:
        p_54_text = find_paragraph_by_text(doc, "model_versions")
    if p_54_text:
        set_paragraph_text(p_54_text,
            "Cada modelo exportado se registra en app.model_versions con: versión semántica "
            "(ej. v1.6.0), is_active (solo una versión activa), model_metadata en JSON con métricas "
            "completas (Log Loss, Brier, ROC-AUC, ECE, feature_columns, trained_at) y description. "
            "Los metadatos también se persisten en ML/models/metadata/vX.X.X_metadata.json. "
            "Para activar una nueva versión en producción se ejecuta: "
            "python scripts/deploy_model.py --version v1.6.0 --activate, "
            "que copia el .joblib, registra en la BD y marca is_active=True. "
            "El historial de versiones disponibles va de v1.0.0 a v2.0.0 (12 versiones)."
        )
        print("[OK] Párrafo 5.4 versionado actualizado")

    # ──────────────────────────────────────────────────────────────
    # 3. SECCIÓN 4 — Scrapping (actualizar comandos)
    # ──────────────────────────────────────────────────────────────
    p_etl_flow = find_paragraph_by_text(doc, "El proceso ETL (Extract, Transform, Load) sigue tres etapas")
    if p_etl_flow:
        set_paragraph_text(p_etl_flow,
            "El proceso ETL sigue tres etapas diferenciadas. En la extracción, los scrapers "
            "visitan páginas de ESPN para obtener boxscores, estadísticas, clasificaciones, lesiones "
            "y cuotas. El script espn/populate_all_games.py realiza la carga histórica inicial completa "
            "de partidos. Para temporadas en curso, scrape_new_boxscores.py realiza actualizaciones "
            "incrementales. Si existen scores faltantes se usan recover_scores.py y "
            "espn/recover_missing_scores.py. audit_full_coverage.py permite auditar la cobertura "
            "del dataset. En la transformación, etl/transform_consolidate.py consolida todas las "
            "fuentes en data/processed/nba_full_dataset.csv. En la carga, load_data.py "
            "inserta los datos en el schema espn de Neon de forma idempotente usando COPY nativo "
            "de PostgreSQL."
        )
        print("[OK] Párrafo 4.2 flujo ETL actualizado")

    # ──────────────────────────────────────────────────────────────
    # 4. SECCIÓN 6 — Backend: actualizar referencia a modelo activo
    # ──────────────────────────────────────────────────────────────
    p_pred_svc = find_paragraph_by_text(doc, "El servicio de predicciones (prediction_service.py)")
    if p_pred_svc:
        set_paragraph_text(p_pred_svc,
            "El servicio prediction_service.py gestiona la carga del modelo ML y la generación de "
            "predicciones. Al iniciar, consulta app.model_versions WHERE is_active=True para determinar "
            "la versión activa (actualmente v1.6.0) y carga el archivo .joblib correspondiente desde "
            "Backend/ml/models/. El resultado es un PredictionResponse con: home_win_probability, "
            "away_win_probability, predicted_home_score, predicted_away_score, predicted_total, "
            "predicted_margin, recommended_bet (home/away/over/under/none), confidence_score y "
            "model_version. Cada predicción genera un snapshot de cuotas en app.odds_snapshots y "
            "un evento en app.outbox. Cuando no hay modelo real disponible, el servicio opera en "
            "modo dummy retornando probabilidades fijas (using_real_predictions=false)."
        )
        print("[OK] Párrafo 6.4 servicio predicciones actualizado")

    # ──────────────────────────────────────────────────────────────
    # Guardar documento
    # ──────────────────────────────────────────────────────────────
    doc.save(str(DOCX_OUT))
    print(f"\n[OK] Documento guardado: {DOCX_OUT}")
    return True


if __name__ == "__main__":
    update_manual()
